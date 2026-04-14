#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档构建器 - Document Builder
基于python-docx构建专业Word文档
"""

import re
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocBuilder:
    """Word文档构建器"""

    def __init__(self, config: Dict = None):
        """
        Args:
            config: 样式配置字典
        """
        self.config = config or self._default_config()
        self.doc = Document()
        self._setup_page()

    def _default_config(self) -> Dict:
        """默认配置"""
        return {
            'fonts': {
                'chinese': '宋体',
                'english': 'Times New Roman',
                'body_size': 11,
                'heading1_size': 18,
                'heading2_size': 14,
                'heading3_size': 12,
            },
            'colors': {
                'primary': (0, 51, 102),
                'secondary': (0, 102, 153),
                'text': (51, 51, 51),
            },
            'spacing': {
                'heading1_before': 24,
                'heading1_after': 12,
                'heading2_before': 18,
                'heading2_after': 8,
                'body_after': 8,
            }
        }

    def _setup_page(self):
        """设置页面布局"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)

    def add_cover(self, data: Dict):
        """添加封面"""
        fonts = self.config['fonts']
        colors = self.config['colors']

        # 顶部空白
        for _ in range(8):
            self.doc.add_paragraph()

        # 主标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(data.get('project_name', '物联网解决方案'))
        run.bold = True
        run.font.name = fonts['chinese']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts['chinese'])
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(*colors['primary'])

        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('技 术 方 案')
        run.bold = True
        run.font.name = fonts['chinese']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts['chinese'])
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(*colors['secondary'])

        # 空白
        for _ in range(6):
            self.doc.add_paragraph()

        # 客户信息
        info_items = [
            ('客户单位', data.get('client_name', '')),
            ('编制日期', data.get('date', '')),
            ('文档版本', data.get('document_version', 'V1.0')),
        ]

        if data.get('contact_person'):
            info_items.append(('联 系 人', data.get('contact_person', '')))
        if data.get('contact_phone'):
            info_items.append(('联系电话', data.get('contact_phone', '')))
        if data.get('contact_email'):
            info_items.append(('联系邮箱', data.get('contact_email', '')))

        for label, value in info_items:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"{label}：{value}")
            run.font.name = fonts['chinese']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts['chinese'])
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(8)

        self.doc.add_page_break()

    def add_toc(self):
        """添加目录"""
        title = self.add_heading1('目 录')

        toc_items = [
            ('1', '项目概述', '3'),
            ('2', '需求分析', '5'),
            ('3', '解决方案设计', '7'),
            ('4', '产品选型清单', '10'),
            ('5', '软件平台介绍', '12'),
            ('6', '实施计划', '14'),
            ('7', '服务与支持', '16'),
            ('8', '附录', '18'),
        ]

        for num, title_text, page in toc_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

            run = p.add_run(f"{num}. {title_text}")
            run.font.name = self.config['fonts']['chinese']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['fonts']['chinese'])
            run.font.size = Pt(12)

            # 添加点线
            dot_run = p.add_run(" " + "." * 50 + " ")
            dot_run.font.size = Pt(12)
            dot_run.font.color.rgb = RGBColor(180, 180, 180)

            page_run = p.add_run(page)
            page_run.font.name = self.config['fonts']['chinese']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['fonts']['chinese'])
            page_run.font.size = Pt(12)

        self.doc.add_page_break()

    def add_heading1(self, text: str) -> any:
        """添加一级标题（使用Word内置"标题1"样式）"""
        p = self.doc.add_heading(text, level=1)
        return p

    def add_heading2(self, text: str):
        """添加二级标题（使用Word内置"标题2"样式）"""
        p = self.doc.add_heading(text, level=2)
        return p

    def add_heading3(self, text: str):
        """添加三级标题（使用Word内置"标题3"样式）"""
        p = self.doc.add_heading(text, level=3)
        return p

    def add_paragraph(self, text: str):
        """添加正文段落"""
        fonts = self.config['fonts']
        spacing = self.config['spacing']

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(spacing['body_after'])
        p.paragraph_format.line_spacing = 1.5

        run = p.add_run(text)
        run.font.name = fonts['chinese']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts['chinese'])
        run.font.size = Pt(fonts['body_size'])
        run.font.color.rgb = RGBColor(*self.config['colors']['text'])

        return p

    def add_bullet(self, text: str, level: int = 0):
        """添加项目符号"""
        fonts = self.config['fonts']

        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
        p.paragraph_format.space_after = Pt(4)

        run = p.add_run(text)
        run.font.name = fonts['chinese']
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fonts['chinese'])
        run.font.size = Pt(fonts['body_size'])
        run.font.color.rgb = RGBColor(*self.config['colors']['text'])

        return p

    def add_table(self, headers: List[str], rows: List[List[str]], col_widths: List[float] = None):
        """添加表格"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # 表头
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header_text

            # 表头样式
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '003366')
            cell._tc.get_or_add_tcPr().append(shading)

            para = cell.paragraphs[0]
            run = para.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = self.config['fonts']['chinese']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['fonts']['chinese'])
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if col_widths and i < len(col_widths):
                cell.width = Cm(col_widths[i])

        # 数据行
        for row_idx, row_data in enumerate(rows):
            row = table.add_row()
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = str(cell_text)

                para = cell.paragraphs[0]
                run = para.runs[0]
                run.font.name = self.config['fonts']['chinese']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['fonts']['chinese'])
                run.font.size = Pt(10)

                # 交替行颜色
                if row_idx % 2 == 1:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F5F5F5')
                    cell._tc.get_or_add_tcPr().append(shading)

        return table

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def add_image(self, image_path: str, alt_text: str = "", width: Inches = None):
        """
        添加图片到文档

        Args:
            image_path: 图片路径（支持绝对路径和相对路径）
            alt_text: 图片描述文字
            width: 图片宽度，默认为页面宽度的80%
        """
        from pathlib import Path

        # 解析图片路径
        img_path = Path(image_path)

        if not img_path.is_absolute():
            # 获取项目根目录 (iot_doc_generator的父目录)
            project_root = Path(__file__).parent.parent

            # 尝试多个可能的路径
            possible_paths = [
                # iot_doc_generator/assets/images/media/image3.png (主要位置)
                project_root / 'assets' / 'images' / 'media' / img_path.name,
                # iot_doc_generator/assets/images/image3.png (如果不在media下)
                project_root / 'assets' / 'images' / img_path.name,
                # 模板目录下的assets/images
                project_root / 'templates' / 'standard' / 'assets' / 'images' / img_path.name,
                # iot_doc_generator根目录下的assets/images
                project_root / image_path,
            ]

            found = False
            for try_path in possible_paths:
                if try_path.exists():
                    img_path = try_path
                    found = True
                    break

            if not found:
                print(f"[WARN] 图片不存在: {image_path}")
                return

        if not img_path.exists():
            print(f"[WARN] 图片不存在: {img_path}")
            return

        # 设置默认宽度为页面宽度的80%
        if width is None:
            width = Inches(5.5)  # A4页面宽度约21cm - 边距，约6.5英寸

        # 添加图片
        try:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=width)

            # 如果有描述文字，添加在图片下方
            if alt_text:
                caption = self.doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"图：{alt_text}")
                caption_run.font.size = Pt(9)
                caption_run.font.color.rgb = RGBColor(128, 128, 128)
                caption_run.italic = True

        except Exception as e:
            print(f"[WARN] 添加图片失败 {img_path}: {e}")

    def render_markdown(self, markdown_text: str):
        """
        渲染Markdown文本到Word

        支持的格式：
        - # 一级标题
        - ## 二级标题
        - ### 三级标题
        - 普通段落
        - - 项目符号
        - | 表格 |
        - ![alt](image_path) 图片
        """
        lines = markdown_text.split('\n')
        in_table = False
        table_headers = []
        table_rows = []

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # 图片处理 ![alt](path)
            import re
            image_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
            image_match = re.match(image_pattern, line)
            if image_match:
                alt_text = image_match.group(1)
                image_path = image_match.group(2)
                self.add_image(image_path, alt_text)
                continue

            # 表格处理
            if '|' in line:
                cells = [c.strip() for c in line.split('|')]
                cells = [c for c in cells if c]  # 移除空单元格

                if cells and not any(c.startswith('---') for c in cells):
                    if not in_table:
                        # 开始表格
                        in_table = True
                        table_headers = cells
                    else:
                        # 继续表格
                        table_rows.append(cells)
                elif in_table and any(c.startswith('---') for c in cells):
                    # 表头分隔线，忽略
                    pass
                elif in_table:
                    # 表格结束，添加表格
                    self.add_table(table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                continue

            if in_table:
                # 表格结束
                self.add_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

            # 标题处理
            if line.startswith('#### '):
                self.add_heading3(line[5:])
            elif line.startswith('### '):
                self.add_heading3(line[4:])
            elif line.startswith('## '):
                self.add_heading2(line[3:])
            elif line.startswith('# '):
                self.add_heading1(line[2:])
            elif line.startswith('- '):
                self.add_bullet(line[2:])
            else:
                # 处理多级列表
                if line.startswith('  - '):
                    self.add_bullet(line[4:], level=1)
                else:
                    self.add_paragraph(line)

        # 处理未结束的表格
        if in_table and table_headers:
            self.add_table(table_headers, table_rows)

    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
        print(f"[OK] 文档已保存: {output_path}")


def test_builder():
    """测试文档构建器"""
    builder = DocBuilder()

    builder.add_cover({
        'project_name': '智能工厂物联网系统',
        'client_name': '广东某某制造有限公司',
        'date': '2026-04-13',
        'document_version': 'V1.0'
    })

    builder.add_toc()

    builder.add_heading1('1. 项目概述')
    builder.add_heading2('1.1 方案概述')
    builder.add_paragraph('这是一段测试文本...')

    builder.add_heading2('1.2 核心价值')
    builder.add_bullet('提升效率')
    builder.add_bullet('降低成本')

    builder.add_heading2('1.3 产品清单')
    builder.add_table(
        ['产品名称', '型号', '数量'],
        [
            ['IoT网关', 'IG200', '10'],
            ['传感器', 'VS-100', '20']
        ]
    )

    output_dir = Path(__file__).parent.parent / 'output'
    output_dir.mkdir(exist_ok=True)
    builder.save(str(output_dir / 'test.docx'))
    print("[OK] 测试文档生成完成")


if __name__ == '__main__':
    test_builder()
