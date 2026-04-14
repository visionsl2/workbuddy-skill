#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word转YAML工具 - Word to YAML Converter
从现有Word文档提取内容，转换为YAML格式，方便迁移和维护

功能：
- 提取标题、段落
- 提取表格数据
- 提取图片引用
- 输出标准YAML格式

使用方法：
    python src/word2yaml.py --input existing.docx --output extracted.yaml
"""

import os
import sys
import argparse
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

import yaml

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("[ERROR] 需要安装 python-docx 库")
    print("运行: pip install python-docx")
    sys.exit(1)


class WordToYamlConverter:
    """Word文档转YAML转换器"""

    def __init__(self):
        self.doc = None
        self.current_section = []
        self.sections = []
        self.tables = []
        self.images = []

    def load_docx(self, file_path: str):
        """加载Word文档"""
        self.doc = Document(file_path)

    def extract_content(self):
        """提取文档内容"""
        if not self.doc:
            return

        # 提取段落
        for para in self.doc.paragraphs:
            self._process_paragraph(para)

        # 提取表格
        for table in self.doc.tables:
            self._process_table(table)

        # 提取图片（内联图片）
        for rel in self.doc.part.rels.values():
            if "image" in rel.reltype:
                self.images.append(rel.target_ref)

    def _process_paragraph(self, para: Paragraph):
        """处理段落"""
        text = para.text.strip()
        if not text:
            return

        # 判断标题级别
        style_name = para.style.name if para.style else ''

        if 'Heading 1' in style_name or style_name == '1':
            self.sections.append({'type': 'heading1', 'text': text})
        elif 'Heading 2' in style_name or style_name == '2':
            self.sections.append({'type': 'heading2', 'text': text})
        elif 'Heading 3' in style_name or style_name == '3':
            self.sections.append({'type': 'heading3', 'text': text})
        else:
            # 普通段落
            self.sections.append({'type': 'paragraph', 'text': text})

    def _process_table(self, table: Table):
        """处理表格"""
        if not table.rows:
            return

        table_data = {
            'headers': [],
            'rows': []
        }

        # 提取表头
        header_cells = table.rows[0].cells
        table_data['headers'] = [cell.text.strip() for cell in header_cells]

        # 提取数据行
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data['rows'].append(row_data)

        self.tables.append(table_data)

    def convert_to_yaml(self) -> Dict:
        """转换为YAML数据结构"""
        result = {
            'meta': {
                'project_name': '',
                'client_name': '',
                'document_version': 'V1.0',
                'date': datetime.now().strftime('%Y-%m-%d'),
            },
            'generated_by': 'word2yaml converter',
            'sections': []
        }

        # 解析章节结构
        current_section = None
        section_content = []

        for item in self.sections:
            if item['type'] == 'heading1':
                # 保存上一个章节
                if current_section:
                    self._finalize_section(result, current_section, section_content)

                # 开始新章节
                current_section = {'title': item['text'], 'subsections': []}
                section_content = []

            elif item['type'] == 'heading2':
                if current_section:
                    # 保存子章节内容
                    if section_content:
                        current_section['subsections'].append({
                            'type': 'content',
                            'text': '\n'.join(section_content)
                        })
                    # 添加子标题
                    current_section['subsections'].append({
                        'type': 'heading2',
                        'text': item['text']
                    })
                    section_content = []

            elif item['type'] == 'paragraph':
                section_content.append(item['text'])

        # 保存最后一个章节
        if current_section:
            self._finalize_section(result, current_section, section_content)

        # 添加表格
        if self.tables:
            result['tables'] = self.tables

        # 添加图片引用
        if self.images:
            result['images'] = self.images

        return result

    def _finalize_section(self, result: Dict, section: Dict, content: List[str]):
        """完成章节处理"""
        section_copy = {
            'title': section['title'],
            'content': '\n'.join(content) if content else ''
        }

        if section.get('subsections'):
            section_copy['subsections'] = section['subsections']

        result['sections'].append(section_copy)

    def save_yaml(self, output_path: str, data: Dict):
        """保存为YAML文件"""
        # 确保目录存在
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            yaml.dump(data, f, allow_unicode=True, default_flow_style=False, sort_keys=False)

        print(f"[OK] YAML已保存: {output_path}")

    def print_summary(self):
        """打印提取摘要"""
        print("\n" + "=" * 50)
        print("提取摘要")
        print("=" * 50)

        # 章节统计
        headings = [s for s in self.sections if 'heading' in s['type']]
        print(f"\n章节数量: {len(headings)}")
        for h in headings[:10]:  # 只显示前10个
            indent = "  " if h['type'] == 'heading2' else ""
            print(f"  {indent}- {h['text'][:50]}")

        if len(headings) > 10:
            print(f"  ... 还有 {len(headings) - 10} 个章节")

        # 表格统计
        print(f"\n表格数量: {len(self.tables)}")
        for i, table in enumerate(self.tables[:5]):
            print(f"  表格{i+1}: {len(table['headers'])}列 x {len(table['rows'])}行")

        # 图片统计
        print(f"\n图片数量: {len(self.images)}")
        for img in self.images[:3]:
            print(f"  - {img}")


def interactive_mode():
    """交互模式"""
    print("\n" + "=" * 50)
    print("Word转YAML 交互模式")
    print("=" * 50)

    result = {}

    # 收集元信息
    print("\n请填写项目信息（直接回车跳过）:")

    result['meta'] = {
        'project_name': input("项目名称: ").strip() or "待命名项目",
        'client_name': input("客户名称: ").strip() or "",
        'document_version': input("文档版本 (默认V1.0): ").strip() or "V1.0",
        'date': datetime.now().strftime('%Y-%m-%d'),
    }

    return result


def main():
    parser = argparse.ArgumentParser(
        description='Word文档转YAML工具 - 从Word文档提取内容并转换为YAML格式',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 基本用法
  python src/word2yaml.py --input 技术方案.docx --output data/projects/extracted.yaml

  # 指定YAML文件格式
  python src/word2yaml.py --input 技术方案.docx --output output.yaml --format yaml

  # 仅预览（不保存）
  python src/word2yaml.py --input 技术方案.docx --preview
        """
    )

    parser.add_argument('--input', '-i', required=True,
                        help='输入Word文档路径')
    parser.add_argument('--output', '-o',
                        help='输出YAML文件路径')
    parser.add_argument('--preview', '-p', action='store_true',
                        help='仅预览，不保存')
    parser.add_argument('--format', choices=['yaml', 'json'], default='yaml',
                        help='输出格式 (默认: yaml)')

    args = parser.parse_args()

    # 检查输入文件
    if not Path(args.input).exists():
        print(f"[ERROR] 文件不存在: {args.input}")
        sys.exit(1)

    # 创建转换器
    converter = WordToYamlConverter()

    # 加载文档
    print(f"\n正在加载: {args.input}")
    try:
        converter.load_docx(args.input)
        print("[OK] 文档加载成功")
    except Exception as e:
        print(f"[ERROR] 加载失败: {e}")
        sys.exit(1)

    # 提取内容
    print("正在提取内容...")
    converter.extract_content()
    print("[OK] 内容提取完成")

    # 打印摘要
    converter.print_summary()

    # 转换为YAML
    data = converter.convert_to_yaml()

    # 输出
    if args.preview:
        print("\n" + "=" * 50)
        print("YAML预览")
        print("=" * 50)
        print(yaml.dump(data, allow_unicode=True, default_flow_style=False))
    else:
        if not args.output:
            # 自动生成输出文件名
            input_path = Path(args.input)
            args.output = str(input_path.with_suffix('.yaml'))

        converter.save_yaml(args.output, data)

        print("\n" + "=" * 50)
        print("转换完成!")
        print("=" * 50)
        print(f"\n提示: 生成的文件可能需要手动调整内容结构，")
        print(f"      以符合模板系统的数据格式要求。")


if __name__ == '__main__':
    main()
